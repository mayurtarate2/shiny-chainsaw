import fitz 

from typing import List, Optional

# document_pipeline/parser.py

import fitz  # PyMuPDF
import docx
import email
import io
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from typing import List, Dict, Union, Optional
import logging
import re
import mimetypes
from pathlib import Path
import zipfile
import xml.etree.ElementTree as ET

logger = logging.getLogger(__name__)

class EnhancedDocumentParser:
    """Enhanced document parser supporting PDF, DOCX, and email formats with improved accuracy."""
    
    def __init__(self):
        self.supported_formats = {'.pdf', '.docx', '.doc', '.eml', '.msg', '.txt'}
    
    def detect_document_type(self, file_path: str) -> str:
        """Detect document type from file extension and content."""
        path = Path(file_path)
        extension = path.suffix.lower()
        
        if extension in self.supported_formats:
            return extension
        
        # Fallback to MIME type detection
        mime_type, _ = mimetypes.guess_type(file_path)
        if mime_type:
            if 'pdf' in mime_type:
                return '.pdf'
            elif 'word' in mime_type or 'officedocument' in mime_type:
                return '.docx'
            elif 'text' in mime_type:
                return '.txt'
        
        raise ValueError(f"Unsupported document format: {extension}")
    
    def extract_text_from_document(self, file_path: str) -> Dict[str, Union[List[str], Dict]]:
        """Extract text from various document formats with enhanced metadata.
        Returns structured data with pages/sections and metadata."""
        doc_type = self.detect_document_type(file_path)
        
        try:
            if doc_type == '.pdf':
                return self._extract_from_pdf(file_path)
            elif doc_type in ['.docx', '.doc']:
                return self._extract_from_docx(file_path)
            elif doc_type in ['.eml', '.msg']:
                return self._extract_from_email(file_path)
            elif doc_type == '.txt':
                return self._extract_from_text(file_path)
            else:
                raise ValueError(f"Unsupported document type: {doc_type}")
                
        except Exception as e:
            logger.error(f"Error extracting text from {doc_type}: {str(e)}")
            raise Exception(f"Failed to extract text from {doc_type}: {str(e)}")
    
    def _extract_from_pdf(self, pdf_path: str) -> Dict[str, Union[List[str], Dict]]:
        """Enhanced PDF extraction with better text quality and metadata."""
        try:
            doc = fitz.open(pdf_path)
            pages = []
            metadata = {
                'total_pages': len(doc),
                'title': doc.metadata.get('title', ''),
                'author': doc.metadata.get('author', ''),
                'subject': doc.metadata.get('subject', ''),
                'format': 'PDF'
            }
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # Try different extraction methods for better quality
                text_blocks = page.get_text("dict")
                text = self._extract_structured_text_from_page(text_blocks)
                
                # Fallback to simple text extraction if structured fails
                if not text.strip():
                    text = page.get_text()
                
                # Clean and enhance the extracted text
                text = self._enhance_text_quality(text)
                pages.append(text)
            
            doc.close()
            return {'pages': pages, 'metadata': metadata}
            
        except Exception as e:
            logger.error(f"Error extracting PDF text: {str(e)}")
            raise
    
    def _extract_structured_text_from_page(self, text_blocks: Dict) -> str:
        """Extract text while preserving structure and formatting."""
        lines = []
        
        for block in text_blocks.get("blocks", []):
            if "lines" in block:
                for line in block["lines"]:
                    line_text = ""
                    for span in line.get("spans", []):
                        text = span.get("text", "")
                        # Preserve important formatting
                        if span.get("flags", 0) & 2**4:  # Bold
                            text = f"**{text}**"
                        line_text += text
                    
                    if line_text.strip():
                        lines.append(line_text.strip())
        
        return "\n".join(lines)
    
    def _extract_from_docx(self, docx_path: str) -> Dict[str, Union[List[str], Dict]]:
        """Enhanced DOCX extraction with structure preservation."""
        try:
            doc = docx.Document(docx_path)
            
            # Extract core properties
            props = doc.core_properties
            metadata = {
                'title': props.title or '',
                'author': props.author or '',
                'subject': props.subject or '',
                'format': 'DOCX',
                'total_paragraphs': len(doc.paragraphs)
            }
            
            sections = []
            current_section = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                
                # Detect headings and create sections
                if self._is_heading(paragraph):
                    if current_section:
                        sections.append('\n'.join(current_section))
                        current_section = []
                    current_section.append(f"## {text}")
                else:
                    current_section.append(text)
            
            # Add final section
            if current_section:
                sections.append('\n'.join(current_section))
            
            # Extract tables
            tables_text = self._extract_tables_from_docx(doc)
            if tables_text:
                sections.append(f"## Tables\n{tables_text}")
            
            return {'pages': sections, 'metadata': metadata}
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {str(e)}")
            raise
    
    def _extract_tables_from_docx(self, doc) -> str:
        """Extract and format table data."""
        tables_text = []
        
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                if any(row_data):  # Skip empty rows
                    table_data.append(" | ".join(row_data))
            
            if table_data:
                tables_text.append("\n".join(table_data))
        
        return "\n".join(tables_text)
    
    def _is_heading(self, paragraph) -> bool:
        """Detect if a paragraph is a heading."""
        # Check if paragraph has heading style
        if paragraph.style.name.startswith('Heading'):
            return True
        
        # Check for bold, larger font, or all caps
        if paragraph.runs:
            first_run = paragraph.runs[0]
            if (first_run.bold or 
                (first_run.font.size and first_run.font.size.pt > 12) or
                paragraph.text.isupper()):
                return True
        
        return False
    
    def _extract_from_email(self, email_path: str) -> Dict[str, Union[List[str], Dict]]:
        """Extract content from email files (.eml, .msg)."""
        try:
            with open(email_path, 'r', encoding='utf-8', errors='ignore') as f:
                email_content = f.read()
            
            msg = email.message_from_string(email_content)
            
            metadata = {
                'subject': msg.get('Subject', ''),
                'from': msg.get('From', ''),
                'to': msg.get('To', ''),
                'date': msg.get('Date', ''),
                'format': 'Email'
            }
            
            # Extract email body
            body_parts = []
            
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        payload = part.get_payload(decode=True)
                        if payload:
                            body_parts.append(payload.decode('utf-8', errors='ignore'))
                    elif part.get_content_type() == "text/html":
                        # Basic HTML to text conversion
                        payload = part.get_payload(decode=True)
                        if payload:
                            html_text = payload.decode('utf-8', errors='ignore')
                            text = self._html_to_text(html_text)
                            body_parts.append(text)
            else:
                payload = msg.get_payload(decode=True)
                if payload:
                    body_parts.append(payload.decode('utf-8', errors='ignore'))
            
            # Structure email content
            email_text = []
            if metadata['subject']:
                email_text.append(f"Subject: {metadata['subject']}")
            if metadata['from']:
                email_text.append(f"From: {metadata['from']}")
            if metadata['to']:
                email_text.append(f"To: {metadata['to']}")
            
            email_text.extend(body_parts)
            
            return {'pages': ['\n'.join(email_text)], 'metadata': metadata}
            
        except Exception as e:
            logger.error(f"Error extracting email text: {str(e)}")
            raise
    
    def _extract_from_text(self, text_path: str) -> Dict[str, Union[List[str], Dict]]:
        """Extract content from plain text files."""
        try:
            with open(text_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            metadata = {
                'format': 'Text',
                'file_size': len(content)
            }
            
            # Split large text files into logical sections
            sections = self._split_text_into_sections(content)
            
            return {'pages': sections, 'metadata': metadata}
            
        except Exception as e:
            logger.error(f"Error extracting text file: {str(e)}")
            raise
    
    def _split_text_into_sections(self, text: str) -> List[str]:
        """Split large text into logical sections."""
        # Split on double newlines or section markers
        sections = re.split(r'\n\s*\n|={3,}|#{2,}', text)
        
        # Filter out empty sections and clean
        cleaned_sections = []
        for section in sections:
            section = section.strip()
            if len(section) > 50:  # Minimum section size
                cleaned_sections.append(section)
        
        return cleaned_sections if cleaned_sections else [text]
    
    def _html_to_text(self, html: str) -> str:
        """Basic HTML to text conversion."""
        # Remove HTML tags
        text = re.sub(r'<[^>]+>', '', html)
        # Clean up whitespace
        text = re.sub(r'\s+', ' ', text)
        return text.strip()
    
    def _enhance_text_quality(self, text: str) -> str:
        """Enhance text quality by cleaning and normalizing."""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Fix common OCR errors
        text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)  # Add space between lowercase and uppercase
        text = re.sub(r'(\w)([.!?])(\w)', r'\1\2 \3', text)  # Add space after punctuation
        
        # Normalize quotes and dashes
        text = text.replace('"', '"').replace('"', '"')
        text = text.replace('', '').replace('', '')
        text = text.replace('—', '-').replace('–', '-')
        
        return text.strip()

# Maintain backward compatibility
def extract_text_from_pdf(pdf_path: str) -> List[str]:
    """Backward compatible function for existing code."""
    parser = EnhancedDocumentParser()
    result = parser.extract_text_from_document(pdf_path)
    return result['pages']


        
        
    